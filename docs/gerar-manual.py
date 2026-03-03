"""
Gera o Manual do Usuário do FinApp em formato Word (.docx)
Fonte: manual-usuario.md (markdown é a fonte única de verdade)
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Tema de cores ──────────────────────────────────────────────────────────────
EMERALD     = RGBColor(0x05, 0x96, 0x69)
EMERALD_HEX = "059669"
EMERALD_BG  = "E6F7F1"
SLATE_900   = RGBColor(0x0F, 0x17, 0x2A)
SLATE_700   = RGBColor(0x33, 0x41, 0x55)
SLATE_500   = RGBColor(0x64, 0x74, 0x8B)
SLATE_200   = "E2E8F0"
SLATE_100   = "F1F5F9"
SLATE_50    = "F8FAFC"
WHITE       = "FFFFFF"
AMBER_BG    = "FFFBEB"
AMBER_BORDER= "F59E0B"

FONT = "Segoe UI"

# ── Helpers de formatação ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/><w:left w:val="none"/>'
        f'<w:bottom w:val="none"/><w:right w:val="none"/>'
        f'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_thin_border_table(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:color="{SLATE_200}"/>'
                f'<w:bottom w:val="single" w:sz="4" w:color="{SLATE_200}"/>'
                f'<w:left w:val="single" w:sz="4" w:color="{SLATE_200}"/>'
                f'<w:right w:val="single" w:sz="4" w:color="{SLATE_200}"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)


def apply_inline_bold(paragraph, text):
    """Parse **bold** markers and add runs with correct formatting."""
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        run.font.name = FONT
        if i % 2 == 1:  # bold segment
            run.bold = True


def set_paragraph_font(paragraph, size=10, color=None, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.bold = bold
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=4, line_spacing=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


# ── Estilos de parágrafo ───────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(22)
    run.font.color.rgb = EMERALD
    run.bold = True
    set_paragraph_spacing(p, before=18, after=4)
    # Underline accent
    p2 = doc.add_paragraph()
    run2 = p2.add_run("─" * 60)
    run2.font.name = FONT
    run2.font.size = Pt(7)
    run2.font.color.rgb = EMERALD
    set_paragraph_spacing(p2, before=0, after=10)
    return p


def add_h2(doc, text):
    # Section number + title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = EMERALD
    run.bold = True
    set_paragraph_spacing(p, before=14, after=3)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = SLATE_700
    run.bold = True
    set_paragraph_spacing(p, before=8, after=3)
    return p


def add_h4(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.color.rgb = SLATE_700
    run.bold = True
    run.italic = True
    set_paragraph_spacing(p, before=6, after=2)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    apply_inline_bold(p, text)
    set_paragraph_font(p, size=10, color=SLATE_900)
    set_paragraph_spacing(p, before=0, after=4)
    return p


def add_tip(doc, text):
    """Blockquote > as a highlighted tip box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, AMBER_BG)
    cell.width = Cm(16)

    # Left border accent via XML
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="16" w:color="{AMBER_BORDER}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    apply_inline_bold(p, text)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(9.5)
        run.font.color.rgb = SLATE_700
        run.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.font.size = Pt(4)
    set_paragraph_spacing(p, before=6, after=6)


# ── Tabelas ────────────────────────────────────────────────────────────────────

def add_table_from_md(doc, rows_data):
    """
    rows_data: list of lists of strings (header row first, no separator row).
    """
    if not rows_data or len(rows_data) < 2:
        return

    col_count = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.clear()

            is_header = (r_idx == 0)
            text = cell_text.strip()

            apply_inline_bold(p, text)

            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(9.5)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    run.font.color.rgb = SLATE_900

            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if is_header:
                set_cell_bg(cell, EMERALD_HEX)
            elif r_idx % 2 == 0:
                set_cell_bg(cell, SLATE_50)
            else:
                set_cell_bg(cell, WHITE)

    set_thin_border_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Parser de markdown ─────────────────────────────────────────────────────────

def parse_markdown(doc, content):
    lines = content.split("\n")
    i = 0
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if table_rows:
            add_table_from_md(doc, table_rows)
        in_table = False
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # Skip YAML front matter / title already handled
        if line.startswith("# FinApp"):
            i += 1
            continue
        if line.startswith("## Manual do Usuário"):
            i += 1
            continue
        if line.startswith("**Versão") or line.startswith("**Data"):
            i += 1
            continue
        if line.strip() == "---":
            if in_table:
                flush_table()
            add_separator(doc)
            i += 1
            continue

        # Sumário — skip
        if line.strip() == "## Sumário":
            while i < len(lines) and not lines[i].startswith("---") and not (lines[i].startswith("## ") and "Sumário" not in lines[i]):
                i += 1
            continue

        # Table row
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator row (|:---:|---|)
            if all(re.match(r'^:?-+:?$', c.strip()) for c in cells if c.strip()):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            add_h4(doc, line[5:].strip())
        elif line.startswith("### "):
            add_h3(doc, line[4:].strip())
        elif line.startswith("## "):
            add_h2(doc, line[3:].strip())
        elif line.startswith("# "):
            add_h1(doc, line[2:].strip())

        # Blockquote / tip
        elif line.startswith("> "):
            text = line[2:].strip()
            # Accumulate multi-line blockquotes
            while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                i += 1
                text += " " + lines[i][2:].strip()
            add_tip(doc, text)

        # Bullet list
        elif line.startswith("- "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            apply_inline_bold(p, text)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(10)
                run.font.color.rgb = SLATE_900
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.6)

        # Numbered list (unlikely in this markdown, handle anyway)
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style="List Number")
            apply_inline_bold(p, text)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(10)
                run.font.color.rgb = SLATE_900

        # Regular paragraph
        else:
            if line.strip():
                add_body(doc, line.strip())

        i += 1

    if in_table:
        flush_table()


# ── Capa ───────────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Spacer
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Logo / título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FinApp")
    run.font.name = FONT
    run.font.size = Pt(42)
    run.font.color.rgb = EMERALD
    run.bold = True
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Manual do Usuário")
    run2.font.name = FONT
    run2.font.size = Pt(18)
    run2.font.color.rgb = SLATE_500
    p2.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Versão 3.0  ·  Março 2026")
    run3.font.name = FONT
    run3.font.size = Pt(11)
    run3.font.color.rgb = SLATE_500
    p3.paragraph_format.space_after = Pt(30)

    doc.add_page_break()


# ── Configuração de página ─────────────────────────────────────────────────────

def configure_page(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("FinApp — Manual do Usuário")
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = SLATE_500

    # Footer with page number
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run()
    run_f.font.name = FONT
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = SLATE_500
    fld = parse_xml(
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:instr=" PAGE "><w:r><w:rPr>'
        '<w:rFonts w:ascii="Segoe UI" w:hAnsi="Segoe UI"/>'
        '<w:sz w:val="16"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    fp._p.append(fld)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path   = os.path.join(script_dir, "manual-usuario.md")
    out_path  = os.path.join(script_dir, "FinApp - Manual do Usuario.docx")

    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    configure_page(doc)
    add_cover(doc)
    parse_markdown(doc, content)

    doc.save(out_path)
    print(f"Manual gerado: {out_path}")


if __name__ == "__main__":
    main()
